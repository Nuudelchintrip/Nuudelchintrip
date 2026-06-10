from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs"
OUT_DIR.mkdir(exist_ok=True)
DOCX_PATH = OUT_DIR / "NuudelchinTrip_Production_Readiness_Blueprint_v1.docx"

BLUE = "2563EB"
DARK = "0F172A"
MUTED = "64748B"
LIGHT_BLUE = "EFF6FF"
LIGHT_GRAY = "F8FAFC"
BORDER = "CBD5E1"
GREEN = "15803D"
GREEN_BG = "F0FDF4"
AMBER = "B45309"
AMBER_BG = "FFFBEB"
RED = "B91C1C"
RED_BG = "FEF2F2"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=140, bottom=100, end=140):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:color"), color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run(run, size=11, color=DARK, bold=False, font="Arial"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold


def add_text(paragraph, text, size=11, color=DARK, bold=False):
    run = paragraph.add_run(text)
    set_run(run, size=size, color=color, bold=bold)
    return run


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    add_text(p, text, size={1: 17, 2: 14, 3: 12}[level], color=BLUE if level < 3 else DARK, bold=True)
    return p


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.22
    if bold_prefix and text.startswith(bold_prefix):
        add_text(p, bold_prefix, bold=True)
        add_text(p, text[len(bold_prefix):])
    else:
        add_text(p, text)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.18
    add_text(p, text)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.18
    add_text(p, text)
    return p


def add_callout(doc, label, text, fill=LIGHT_BLUE, label_color=BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.5)
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=180, start=220, bottom=180, end=220)
    set_table_borders(table, color=fill, size="2")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    add_text(p, label.upper(), size=9, color=label_color, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.18
    add_text(p2, text, size=11, color=DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_status_table(doc):
    data = [
        ("Public website", "Ажиллаж байна", "Нүүр, танилцуулга, бүртгэл, mobile public layout"),
        ("Auth ба role routing", "Хэсэгчлэн", "Supabase auth ажиллаж байгаа ч production OTP/SMTP дутуу"),
        ("Жолоочийн чиглэл", "Хэсэгчлэн", "Чиглэл үүсгэх бодит; verification document upload дутуу"),
        ("Аялагчийн захиалга", "Хэсэгчлэн", "Хайлт, суудал сонголт, booking бий; lifecycle бүрэн биш"),
        ("Төлбөр", "Хэсэгчлэн", "Баримт upload/admin approve бий; settlement/refund/atomic flow дутуу"),
        ("Аяллын явц", "Дутуу", "Эхлүүлэх, дуусгах, audit log, notification бүрэн холбогдоогүй"),
        ("Дайвар ачаа", "Хэсэгчлэн", "Request бий; detail, proof, delivery code lifecycle дутуу"),
        ("Admin", "Хэсэгчлэн", "Зарим queue бодит; reports/logs/operations бүрэн биш"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [Inches(1.55), Inches(1.15), Inches(3.8)]
    headers = ["Хэсэг", "Төлөв", "Тайлбар"]
    for i, (cell, width, text) in enumerate(zip(table.rows[0].cells, widths, headers)):
        cell.width = width
        set_cell_shading(cell, BLUE)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        add_text(p, text, size=10, color=WHITE, bold=True)
    set_repeat_table_header(table.rows[0])
    for section, status, note in data:
        row = table.add_row()
        for i, (cell, width, text) in enumerate(zip(row.cells, widths, (section, status, note))):
            cell.width = width
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if i == 1:
                fill, color = (
                    (GREEN_BG, GREEN) if status == "Ажиллаж байна"
                    else (AMBER_BG, AMBER) if status == "Хэсэгчлэн"
                    else (RED_BG, RED)
                )
                set_cell_shading(cell, fill)
                add_text(cell.paragraphs[0], text, size=9.5, color=color, bold=True)
            else:
                add_text(cell.paragraphs[0], text, size=9.5)
    set_table_borders(table)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_phase_table(doc):
    rows = [
        ("1", "Auth, OTP, SMTP", "Production SMS/OTP, email confirmation, reset password", "Хэрэглэгч найдвартай бүртгүүлж нэвтэрнэ"),
        ("2", "Жолоочийн verification", "Document upload, admin review, approve/reject", "Approved жолооч л чиглэл нийтэлнэ"),
        ("3", "Аялагчийн аялал", "Миний аялал, booking list/detail, зөв empty state", "Аялагч өөрийн захиалгаа бүрэн харна"),
        ("4", "Суудлын удирдлага", "Seat hold, expiry, reject/cancel release", "Суудал давхардахгүй, түгжигдэхгүй"),
        ("5", "Booking lifecycle", "Role-validated RPC state transitions", "Төлөв зөв дарааллаар, audit-тай шилжинэ"),
        ("6", "Төлбөр", "Admin account, proof review, transaction, refund/payout", "Мөнгөний урсгал хяналттай болно"),
        ("7", "Аялал эхлэх/дуусах", "Driver actions, timestamps, logs, notifications", "Бодит аяллын lifecycle дуусна"),
        ("8", "Мэдэгдэл ба лог", "In-app notification, admin alerts, audit history", "Хэрэглэгч дараагийн алхмаа мэднэ"),
        ("9", "Дайвар ачаа", "Cargo detail, proof, delivery code, status lifecycle", "Route-based cargo add-on бүтэн ажиллана"),
        ("10", "Admin operations", "Users, reports, payments, routes, disputes real data", "Админ өдөр тутмын ажиллагааг удирдана"),
        ("11", "Mobile, security, QA", "Responsive QA, RLS audit, tests, monitoring", "Production эрсдэл буурна"),
        ("12", "Launch ба handoff", "Policies, backups, runbook, customer acceptance", "Нээлттэй ашиглалтад шилжинэ"),
    ]
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [Inches(0.42), Inches(1.35), Inches(2.58), Inches(2.15)]
    headers = ["#", "Үе шат", "Хийх ажил", "Дууссан шалгуур"]
    for cell, width, text in zip(table.rows[0].cells, widths, headers):
        cell.width = width
        set_cell_shading(cell, BLUE)
        set_cell_margins(cell, top=110, bottom=110)
        add_text(cell.paragraphs[0], text, size=9, color=WHITE, bold=True)
    set_repeat_table_header(table.rows[0])
    for phase, title, work, done in rows:
        row = table.add_row()
        for idx, (cell, width, text) in enumerate(zip(row.cells, widths, (phase, title, work, done))):
            cell.width = width
            set_cell_margins(cell, top=90, bottom=90)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if idx == 0:
                set_cell_shading(cell, LIGHT_BLUE)
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_text(p, text, size=9, color=BLUE, bold=True)
            else:
                add_text(cell.paragraphs[0], text, size=8.8, bold=(idx == 1))
    set_table_borders(table)


def add_acceptance_section(doc, title, items):
    add_heading(doc, title, 2)
    for item in items:
        add_bullet(doc, f"☐ {item}")


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.78)
section.bottom_margin = Inches(0.72)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.35)
section.footer_distance = Inches(0.35)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Arial"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor.from_string(DARK)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.22

for level, size, before, after, color in (
    (1, 17, 18, 10, BLUE),
    (2, 14, 14, 7, BLUE),
    (3, 12, 10, 5, DARK),
):
    style = styles[f"Heading {level}"]
    style.font.name = "Arial"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

header = section.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
add_text(hp, "NuudelchinTrip | Production Readiness Blueprint", size=8.5, color=MUTED)

footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_text(fp, "NuudelchinTrip • 2026", size=8.5, color=MUTED)

# Cover
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(34)
p.paragraph_format.space_after = Pt(8)
add_text(p, "NUUDELCHINTRIP", size=12, color=BLUE, bold=True)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(10)
add_text(p, "Production Readiness\nBlueprint", size=30, color=DARK, bold=True)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(28)
add_text(
    p,
    "Ажилладаг MVP-ээс олон нийт ашиглах найдвартай платформ руу шилжих хэрэгжүүлэлтийн зураглал",
    size=14,
    color=MUTED,
)

add_callout(
    doc,
    "Одоогийн үнэлгээ",
    "NuudelchinTrip нь үндсэн Supabase холболттой, ажилладаг MVP/demo болсон. "
    "Гэхдээ production хэрэглээнд нээхийн өмнө authentication, verification, суудлын удирдлага, "
    "төлбөр, аяллын lifecycle, мэдэгдэл, security болон QA-г бүрэн дуусгах шаардлагатай.",
)

meta = doc.add_table(rows=4, cols=2)
meta.alignment = WD_TABLE_ALIGNMENT.LEFT
meta.autofit = False
for row, (label, value) in zip(
    meta.rows,
    (
        ("Баримтын хувилбар", "v1.0"),
        ("Шалгасан огноо", "2026 оны 6 сарын 7"),
        ("Live website", "https://nuudelchintrip.com"),
        ("Үндсэн бүтээгдэхүүн", "Аялагч ба жолоочийг чиглэлээр холбох платформ"),
    ),
):
    row.cells[0].width = Inches(1.65)
    row.cells[1].width = Inches(4.85)
    for cell in row.cells:
        set_cell_margins(cell, top=85, bottom=85)
    set_cell_shading(row.cells[0], LIGHT_GRAY)
    add_text(row.cells[0].paragraphs[0], label, size=9.5, color=MUTED, bold=True)
    add_text(row.cells[1].paragraphs[0], value, size=9.5)
set_table_borders(meta, color=BORDER)

doc.add_page_break()

add_heading(doc, "1. Blueprint-ийн зорилго", 1)
add_body(
    doc,
    "Энэ баримт нь NuudelchinTrip-ийг demo түвшнээс бодит хэрэглэгч, жолооч, админ ашиглах production платформ болгоход "
    "шаардлагатай ажлыг эрэмбэлж, шат бүрийн дууссан шалгуурыг тогтооно."
)
add_body(
    doc,
    "Бүтээгдэхүүний гол урсгал нь аялагч жолоочийн нийтэлсэн чиглэлийг хайж, суудлаа сонгон захиалах явдал. "
    "Дайвар ачаа нь жолоочийн чиглэл дээр суурилсан нэмэлт боломж байна."
)

add_callout(
    doc,
    "Launch шийдвэр",
    "Одоогийн хувилбарыг захиалагчид MVP/demo байдлаар үзүүлж болно. Харин танихгүй олон хэрэглэгчид "
    "нээлттэй ашиглуулахын өмнө энэхүү blueprint-ийн P0 ажлууд заавал дууссан байна.",
    fill=AMBER_BG,
    label_color=AMBER,
)

add_heading(doc, "2. Одоогийн бодит төлөв", 1)
add_status_table(doc)

add_heading(doc, "3. Production launch blocker", 1)
for title, description in (
    ("Production OTP ба email", "Demo баталгаажуулалтыг жинхэнэ SMS/OTP болгох; custom SMTP, password reset, rate limit нэмэх."),
    ("Жолоочийн бичиг баримт", "Үнэмлэх, машины гэрчилгээ, зураг бодитоор upload хийж, админ шалгах."),
    ("Суудлын түгжээ ба суллалт", "Pending хүсэлтэд хугацаатай hold хийх; reject/cancel/expiry үед автоматаар суллах."),
    ("Booking төлөвийн хамгаалалт", "Төлөв шилжилтийг client update биш role-validated RPC/серверийн функцээр хийх."),
    ("Төлбөрийн найдвартай урсгал", "Баримт, approval, booking update-ийг transaction болгох; refund/payout дүрэмтэй байх."),
    ("Аялал эхлүүлэх ба дуусгах", "Жолооч action, timestamp, status log, аялагч/админ мэдэгдлийг холбох."),
    ("Fake/placeholder UX арилгах", "Ажиллахгүй чат, report, түр код, mock хүн/утас/данс, placeholder мэдээллийг бүрэн цэвэрлэх."),
    ("Security ба RLS", "Role, verification, phone_verified зэрэг sensitive field-ийг хэрэглэгч өөрөө өөрчилж чадахгүй болгох."),
):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    add_text(p, f"{title}: ", bold=True, color=DARK)
    add_text(p, description, color=MUTED)

doc.add_page_break()

add_heading(doc, "4. Хэрэгжүүлэх 12 үе шат", 1)
add_body(doc, "Ажлыг доорх дарааллаар хийх нь өгөгдөл, эрх, төлөвийн зөрчил үүсэх эрсдэлийг хамгийн бага байлгана.")
add_phase_table(doc)

doc.add_page_break()

add_heading(doc, "5. Үндсэн хэрэглэгчийн урсгал", 1)
add_heading(doc, "5.1 Аялагч", 2)
for step in (
    "Бүртгүүлэх → утас баталгаажуулах → аялагчийн мэдээллээ бүрдүүлэх",
    "Жолоочийн чиглэл хайх → чиглэлийн дэлгэрэнгүй үзэх",
    "Сул суудлаас өөрөө сонгох → захиалгын хүсэлт илгээх",
    "Жолооч зөвшөөрөх → төлбөрийн баримт оруулах",
    "Админ төлбөр батлах → аялал баталгаажих",
    "Аялал эхлэх/дуусах мэдэгдэл авах → үнэлгээ өгөх",
):
    add_number(doc, step)

add_heading(doc, "5.2 Жолооч", 2)
for step in (
    "Бүртгүүлэх → утас баталгаажуулах → машин, үнэмлэхийн мэдээлэл оруулах",
    "Админ verification approve хийх",
    "Чиглэл, огноо, цаг, үнэ, суудлын байрлал нийтлэх",
    "Ирсэн суудлын болон дайвар ачааны хүсэлтийг зөвшөөрөх/татгалзах",
    "Аяллыг эхлүүлэх → дуусгах",
    "Орлого, төлбөр, үнэлгээний түүхээ харах",
):
    add_number(doc, step)

add_heading(doc, "5.3 Админ", 2)
for step in (
    "Жолоочийн бичиг баримт шалгах",
    "Төлбөрийн баримт батлах эсвэл татгалзах",
    "Аялал эхэлсэн/дууссан болон маргааны мэдэгдэл хүлээн авах",
    "Хэрэглэгч, чиглэл, захиалга, ачаа, report-ыг хянах",
    "Refund, suspension, dispute зэрэг ажиллагааг audit log-той хийх",
):
    add_number(doc, step)

add_heading(doc, "6. Database ба backend шаардлага", 1)
for item in (
    "profiles: role болон verification field-үүдийг зөвхөн хамгаалагдсан RPC/admin update хийнэ.",
    "driver_profiles: document path, verification reviewer, reviewed_at, rejection_reason хадгална.",
    "trips: structured location, seats, status, cargo permission болон expiry дүрэмтэй байна.",
    "passenger_bookings: selected seats, hold_expires_at, state transition timestamps хадгална.",
    "payments/proofs: booking/cargo reference, amount, reviewer, reviewed_at, rejection reason хадгална.",
    "trip_status_logs: төлөв бүрийн өмнөх/дараах утга, actor, timestamp, note хадгална.",
    "notifications: recipient, event type, read_at, deeplink хадгална.",
    "reviews/reports: зөвхөн холбоотой, дууссан transaction-аас үүсэх серверийн шалгалттай байна.",
):
    add_bullet(doc, item)

add_callout(
    doc,
    "Чухал дүрэм",
    "Booking, payment approval, trip start/end, seat release зэрэг олон хүснэгт өөрчилдөг үйлдлийг "
    "frontend-ээс дараалсан update хэлбэрээр бус database transaction/RPC эсвэл server function-оор гүйцэтгэнэ.",
    fill=RED_BG,
    label_color=RED,
)

doc.add_page_break()

add_heading(doc, "7. Security checklist", 1)
security_items = [
    "Admin route нь UI нууснаар бус database role-оор хамгаалагдсан",
    "User role, phone_verified, is_suspended, verification_status-ийг өөрөө update хийж чадахгүй",
    "Storage upload file type, хэмжээ, owner path, signed URL бодлоготой",
    "OTP resend, login, register, support form дээр rate limit тавьсан",
    "Booking/cargo status transition бүр role болон одоогийн төлвийг шалгадаг",
    "Payment approval, verification, suspension бүр audit log үүсгэдэг",
    "Supabase service role key frontend bundle-д байхгүй",
    "Production environment variables зөвхөн Vercel/Supabase secret store-д хадгалагдсан",
    "Backup, restore test, security alert болон access review төлөвлөгөөтэй",
]
for item in security_items:
    add_bullet(doc, f"☐ {item}")

add_heading(doc, "8. UI/UX ба content checklist", 1)
for item in (
    "Customer UI бүхэлдээ ойлгомжтой Монгол хэлтэй; role, admin, review зэрэг хөгжүүлэлтийн үг хэрэглээгүй",
    "Аялагчийн самбарт жолоочийн action харагдахгүй; жолоочийн самбарт аялагчийн action харагдахгүй",
    "Хуудас бүр дараагийн хийх үйлдлийг нэг үндсэн CTA-аар тодорхой харуулдаг",
    "Ажиллахгүй товч, mock тоо, fake review, fake банкны данс, placeholder хүн байхгүй",
    "Empty, loading, error, success, offline state бүр хэрэглэгчид тайлбар өгдөг",
    "Mobile дээр sidebar drawer, form, seat picker, modal, table хэвтээ overflow-гүй",
    "Нэг хуудас нэг H1, form бүр label, autocomplete, keyboard focus, contrast шаардлага хангадаг",
    "SEO title, description, Open Graph Монгол UTF-8 текст зөв харагддаг",
):
    add_bullet(doc, f"☐ {item}")

add_heading(doc, "9. QA test matrix", 1)
qa = [
    ("Аялагч", "Register → OTP → хайлт → суудал → booking → payment → trip → review"),
    ("Жолооч", "Register → document verification → route → accept → start/end → earnings"),
    ("Ачаа илгээгч", "Cargo-enabled route → request → payment → pickup → delivery code"),
    ("Админ", "Verification → payment → reports → suspension/refund → audit log"),
    ("Security", "Role bypass, RLS, duplicate seat, invalid status, file upload abuse"),
    ("Mobile", "360px, 390px, tablet, desktop дээр бүх core flow"),
]
table = doc.add_table(rows=1, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False
for cell, width, text in zip(table.rows[0].cells, [Inches(1.35), Inches(5.15)], ["Тест", "Заавал шалгах урсгал"]):
    cell.width = width
    set_cell_shading(cell, BLUE)
    set_cell_margins(cell)
    add_text(cell.paragraphs[0], text, size=9.5, color=WHITE, bold=True)
for label, flow in qa:
    row = table.add_row()
    row.cells[0].width = Inches(1.35)
    row.cells[1].width = Inches(5.15)
    for cell in row.cells:
        set_cell_margins(cell)
    add_text(row.cells[0].paragraphs[0], label, size=9.5, bold=True)
    add_text(row.cells[1].paragraphs[0], flow, size=9.5)
set_table_borders(table)

doc.add_page_break()

add_heading(doc, "10. Phase acceptance criteria", 1)
add_acceptance_section(
    doc,
    "P0 — Олон нийтэд нээхээс өмнө",
    [
        "Шинэ traveler, driver, cargo account бүртгүүлж, OTP баталгаажуулж чаддаг",
        "Driver document upload ба admin approval бодитоор ажилладаг",
        "Seat hold/release болон booking state machine давхар захиалга үүсгэдэггүй",
        "Төлбөрийн баримт, approval, booking confirmation transaction байдлаар ажилладаг",
        "Driver аялал эхлүүлж/дуусгаж, аялагч болон админ мэдэгдэл авдаг",
        "Core flow-д mock data, ажиллахгүй товч, placeholder action үлдээгүй",
        "RLS/security test болон 4 role-ийн end-to-end test амжилттай",
    ],
)
add_acceptance_section(
    doc,
    "P1 — Захиалагчид production байдлаар хүлээлгэн өгөх",
    [
        "Cargo lifecycle, review, report, support бүр real data-тай",
        "Terms, privacy, cancellation, refund, prohibited cargo policy батлагдсан",
        "Monitoring, error tracking, backups, alert тохирсон",
        "Mobile болон accessibility QA дууссан",
        "Admin runbook, deployment guide, test accounts, handoff documentation бэлэн",
    ],
)

add_heading(doc, "11. Хүлээлгэн өгөх багц", 1)
for item in (
    "Production source code ба version tag",
    "Supabase schema/migration болон RLS policy",
    "Vercel environment variable жагсаалт",
    "Admin хэрэглэгч үүсгэх заавар",
    "Operational runbook: payment, verification, report, refund",
    "QA report болон үлдсэн known issue жагсаалт",
    "Backup/restore болон incident response товч заавар",
):
    add_bullet(doc, item)

add_callout(
    doc,
    "Эцсийн дүгнэлт",
    "NuudelchinTrip-ийн үндсэн концепц зөв бөгөөд Supabase-тэй ажилладаг суурь бий. "
    "Дараагийн ажил нь шинэ feature олноор нэмэх бус, дээрх 12 үе шатыг дарааллаар дуусгаж "
    "core passenger-driver booking flow-ийг найдвартай, хамгаалагдсан, хэмжигдэхүйц болгох юм.",
)

doc.save(DOCX_PATH)
print(DOCX_PATH)
